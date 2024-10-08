from django.shortcuts import render, redirect, get_object_or_404
from django.core.files.storage import FileSystemStorage
from django.contrib import messages
import pandas as pd
from .models import QuestionBank,InputSuggestion,InputSuggestionImage, InputSuggestionDocument, ExamName, Subject, Area, PartName, ChapterName,TopicName, QuoteIdiomPhrase
from django.db.models import Max, Count, Value, Q
from .forms import UploadFileForm, InputSuggestionForm
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image as PILImage
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from django.http import FileResponse, HttpResponse
from PIL import Image as PILImage
from django.http import JsonResponse
from django.contrib.auth.decorators import login_required
import json
from django.utils import timezone
from datetime import timedelta
from django.db import transaction
from accounts.models import User
import csv
from io import BytesIO
from django.http import HttpResponse
from .models import Report
from django.template.loader import render_to_string
from django.db.models.functions import Concat
from django.contrib.auth import get_user_model
from django.conf import settings

User = get_user_model()

# ************************* Generate Test Word file Start *********************************************

def clean_text(text):
    """Utility function to clean and format text for the document."""
    return text.strip() if text else ''


def generate_questions(request):
    try:
        # Create an in-memory file object
        buffer = BytesIO()

        # Setup document file to save generated word content
        today = datetime.today().strftime('%Y-%m-%d')
        file_name = f'all_questions_{today}.docx'
        document = Document()

        # Add content to the document
        for question in QuestionBank.objects.all():
            if question.question_sub_type == 'simple_type':
                add_simple_type(question, document)
            elif question.question_sub_type == 'r_and_a_type':
                add_r_and_a_type(question, document)
            elif question.question_sub_type == 'list_type_1':
                add_list_type_1(question, document)
            elif question.question_sub_type == 'list_type_2':
                add_list_type_2(question, document)
            
            # Add a space between questions
            document.add_paragraph()

        # Save the document to the in-memory file object
        document.save(buffer)
        buffer.seek(0)

        # Return the generated file as a downloadable response
        response = FileResponse(buffer, as_attachment=True, filename=file_name)
        response['Content-Disposition'] = f'attachment; filename={file_name}'
        return response

    except Exception as e:
        return HttpResponse(f"An error occurred: {str(e)}", status=500)


def add_simple_type(question, document):
    """Add simple type question to the document."""
    document.add_paragraph(f"({question.question_number}). {clean_text(question.question_part)}")
    add_options_and_answers(document, question)

def add_r_and_a_type(question, document):
    """Add reason and assertion type question to the document."""
    document.add_paragraph(f"({question.question_number}). {clean_text(question.question_part_first)}")
    document.add_paragraph(f"{clean_text(question.question_part_third)}")
    add_options_and_answers(document, question)

def add_list_type_1(question, document):
    """Add list type 1 question to the document."""
    document.add_paragraph(f"({question.question_number}). {clean_text(question.question_part_first)}")
    for i in range(1, 9):  # Adjusted to match the model's 8 list rows
        list_row = getattr(question, f'list_1_row{i}', None)
        if list_row:
            document.add_paragraph(f"{i}. {clean_text(list_row)}")
    document.add_paragraph(f"{clean_text(question.question_part_third)}")
    add_options_and_answers(document, question)

def add_list_type_2(question, document):
    """Add list type 2 question to the document."""
    document.add_paragraph(f"({question.question_number}). {clean_text(question.question_part_first)}")
    
    # Create table with two columns
    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "LIST - I"
    hdr_cells[1].text = "LIST - II"
    
    # Add the list names, if available
    if question.list_1_name:
        hdr_cells[0].text += f"\n({clean_text(question.list_1_name)})"
    if question.list_2_name:
        hdr_cells[1].text += f"\n({clean_text(question.list_2_name)})"

    # Add list rows
    labels = 'ABCDEFGHIJKLMNOPQRSTUVWXYZ'
    for i in range(1, 9):  # Adjusted to match the model's 8 list rows
        list_1_item = clean_text(getattr(question, f'list_1_row{i}', ''))
        list_2_item = clean_text(getattr(question, f'list_2_row{i}', ''))

        if list_1_item or list_2_item:
            row = table.add_row().cells
            row[0].text = f"{labels[i-1]}. {list_1_item}" if list_1_item else ''
            row[1].text = f"{i}. {list_2_item}" if list_2_item else ''

    # Add question details and options below the table
    if question.question_part_third:
        document.add_paragraph(f"{clean_text(question.question_part_third)}")
    
    add_options_and_answers(document, question)

def add_options_and_answers(document, question):
    """Add options and answers to the document."""
    for opt in ['a', 'b', 'c', 'd']:
        option_text = getattr(question, f'answer_option_{opt}', None)
        if option_text:
            document.add_paragraph(f"({opt.lower()}) {clean_text(option_text)}")
    
    document.add_paragraph(f"Correct Answer: {clean_text(question.correct_answer_choice)}")
    document.add_paragraph(f"Solution: {clean_text(question.correct_answer_description)}")
    
    # Format created_at datetime
    created_at_str = question.created_at.strftime('%Y-%m-%d %H:%M:%S')
    document.add_paragraph(f"Created At: {created_at_str}")
    
    # Handle created_by (User object from Django built-in User model)
    if question.created_by:
        created_by_str = question.created_by.get_full_name() or question.created_by.username
        document.add_paragraph(f"Created By: {clean_text(created_by_str)}")


# ************************* Generate Test Word file End *********************************************


# ************************* Generate Clas Plus Word file Start *********************************************

import re

def clean_text(text):
    """Utility function to clean and format text by removing extra newlines and spaces."""
    if not text:
        return ''
    # Strip leading and trailing spaces
    text = text.strip()
    # Replace multiple newlines or newline + spaces with a single space
    text = re.sub(r'\s*\n\s*', ' ', text)
    # Ensure no multiple spaces remain after replacements
    text = re.sub(r'\s+', ' ', text)
    return text



def set_no_border(cell):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def generate_questions_document(request):
    try:
        # Setup directory and document file to save generated Word file
        base_dir = os.path.join(settings.MEDIA_ROOT, 'word_file')
        os.makedirs(base_dir, exist_ok=True)
        today = datetime.today().strftime('%Y-%m-%d')
        file_name = f'class_plus_questions_{today}.docx'
        file_path = os.path.join(base_dir, file_name)
        document = Document()

        for question in QuestionBank.objects.all():
            # Create main table for the question and set styles
            table = document.add_table(rows=0, cols=3)
            table.style = 'Table Grid'

            # Construct the question text
            # Construct the question text based on question_sub_type
            if question.question_sub_type == "r_and_a_type":
                question_text = (
                    clean_text(question.question_part_first or '') + "\n" +
                    (clean_text(question.assertion) if question.assertion else '') + "\n" +
                    (clean_text(question.reason) if question.reason else '') + "\n" +
                    clean_text(question.question_part_third or '')
                )
            else:
                # Default construction for other question types
                if question.question_part and question.question_part.strip():
                    question_text = clean_text(question.question_part)
                else:
                    question_text = (
                        clean_text(question.question_part_first or '') + "\n" +
                        ("Assertion (A): " + clean_text(question.assertion) if question.assertion else '') + "\n" +
                        ("Reason (R): " + clean_text(question.reason) if question.reason else '') + "\n" +
                        clean_text(question.question_part_third or '')
                    )


            # Add question info to the main table
            q_row = table.add_row().cells
            q_row[0].text = 'Question'

            if question.question_sub_type == 'list_type_1':
                # Handle list_i_selection_text or list-I type questions
                q_row[1].text = f"{clean_text((question.question_part_first) or '')}\n"
                options_format = "\n".join([
                    f"{i}. {getattr(question, f'list_1_row{i}', '')}" for i in range(1, 9)
                    if getattr(question, f'list_1_row{i}', '')
                ])
                q_row[1].text += options_format  + "\n" + clean_text((question.question_part_third) or '')
            # Handle list type questions with sub-tables
            elif question.list_1_name and question.list_2_name:
                sub_table = document.add_table(rows=1, cols=2)
                sub_table.style = 'Table Grid'
                sub_hdr_cells = sub_table.rows[0].cells
                sub_hdr_cells[0].text = f"LIST - I\n({clean_text(question.list_1_name)})"
                sub_hdr_cells[1].text = f"LIST - II\n({clean_text(question.list_2_name)})"

                for i in range(1, 9):
                    list_1_option = getattr(question, f'list_1_row{i}', '')
                    list_2_option = getattr(question, f'list_2_row{i}', '')
                    if not list_2_option:
                        break  # Stop loop if list_2_item is null
                    row_cells = sub_table.add_row().cells
                    row_cells[0].text = f"{chr(64+i)}. {list_1_option}"
                    row_cells[1].text = f"{i}. {list_2_option}"

                q_row[1]._element.clear_content()
                p = q_row[1].paragraphs[0] if q_row[1].paragraphs else q_row[1].add_paragraph()
                p.add_run((clean_text(question.question_part_first) or '') + "\n")
                q_row[1]._element.append(sub_table._element)

                # Add 'Codes:' text below the sub-table within the same cell
                p = q_row[1].add_paragraph()
                p.add_run("\nCodes:\t A\t B\t C\t D")
            else:
                q_row[1].text = question_text

            # Merging cells for question text and image
            q_row[1].merge(q_row[2])

            # Handling image insertion if available
            if question.image:
                image_path = question.image.path
                pil_img = PILImage.open(image_path)

                # Check if the image is in 'RGBA' mode (i.e., if it's a PNG with transparency)
                if pil_img.mode == 'RGBA':
                    pil_img = pil_img.convert('RGB')  # Convert to 'RGB' mode to discard transparency

                # Save the image in JPEG format regardless of the original format
                img_io = BytesIO()
                pil_img.save(img_io, format='JPEG')
                img_io.seek(0)

                # Add the image to the document between question_part_first and question_part_third
                paragraph = q_row[1].add_paragraph()
                run = paragraph.add_run()
                run.add_picture(img_io, width=Inches(1.5))

            # Table-type questions
            table_heads = ['table_head_a', 'table_head_b', 'table_head_c', 'table_head_d']
            data_fields = [
                [getattr(question, f'head_a_data{j}', None) for j in range(1, 5)],
                [getattr(question, f'head_b_data{j}', None) for j in range(1, 5)],
                [getattr(question, f'head_c_data{j}', None) for j in range(1, 5)],
                [getattr(question, f'head_d_data{j}', None) for j in range(1, 5)]
            ]

            filtered_heads_data = [
                (head, [data for data in datas if data])
                for head, datas in zip(table_heads, data_fields)
                if getattr(question, head, None) and any(datas)
            ]

            if filtered_heads_data:
                sub_table = document.add_table(rows=len(filtered_heads_data[0][1]) + 1, cols=len(filtered_heads_data))
                sub_table.style = 'Table Grid'
                hdr_cells = sub_table.rows[0].cells
                for idx, (head, _) in enumerate(filtered_heads_data):
                    hdr_cells[idx].text = getattr(question, head, "")
                for col_idx, (head, data_list) in enumerate(filtered_heads_data):
                    for row_idx, data in enumerate(data_list):
                        cell = sub_table.cell(row_idx + 1, col_idx)
                        cell.text = data
                q_row[1]._element.append(sub_table._element)

            # Adding the Type row based on conditions
            type_row = table.add_row().cells
            type_row[0].text = 'Type'

            # Set the type text based on the question_sub_type
            if question.question_sub_type == 'true_and_false_type':
                type_row[1].text = 'true_false'
            elif question.question_sub_type == 'fill_in_the_blank_type':
                type_row[1].text = 'fill_ups'
            else:
                type_row[1].text = 'multiple_choice'

            # Merging cells for consistent formatting
            type_row[1].merge(type_row[2])

            # Adding options (including for true/false and fill-ups)
            valid_options = ['a', 'b', 'c', 'd']
            correct_answer = question.correct_answer_choice.lower() if question.correct_answer_choice else None
            for opt in valid_options:
                option_text = getattr(question, f"answer_option_{opt}", None)
                if option_text:
                    opt_row = table.add_row().cells
                    opt_row[0].text = 'Option'
                    opt_row[1].text = f"{opt.upper()}. {option_text}"
                    opt_row[2].text = 'correct' if opt == correct_answer else 'incorrect'

            if question.question_sub_type == 'fill_in_the_blank_type':
                option_row = table.add_row().cells
                option_row[0].text = 'Option'
                option_row[1].text = question.correct_answer_choice or '___'

            if question.question_sub_type == 'true_and_false_type':
                answer_row = table.add_row().cells
                answer_row[0].text = 'Answer'

                # Get the correct answer letter (e.g., 'a', 'b') and fetch the corresponding option text
                correct_answer_letter = question.correct_answer_choice.lower()  # Get the correct answer letter in lowercase
                correct_answer_text = getattr(question, f"answer_option_{correct_answer_letter}", None)  # Fetch the text of the correct answer option

                if correct_answer_text:
                    answer_row[1].text = correct_answer_text  # Set the correct answer option text
                else:
                    answer_row[1].text = "Answer not available"  # Fallback in case the option text is missing
                
                answer_row[1].merge(answer_row[2])

            solution_row = table.add_row().cells
            solution_row[0].text = 'Solution'
            solution_row[1].text = clean_text(question.correct_answer_description)
            solution_row[1].merge(solution_row[2])

            marks_row = table.add_row().cells
            marks_row[0].text = 'Marks'
            marks_row[1].text = str(question.marks)
            marks_row[2].text = str(question.negative_marks)

            document.add_paragraph()  # Add space between questions

        document.save(file_path)
        
        # Return the generated file as a downloadable response
        response = FileResponse(open(file_path, 'rb'), as_attachment=True, filename=file_name)
        return response

    except Exception as e:
        return HttpResponse(f"An error occurred: {str(e)}", status=500)


# ************************* Generate Clas Plus Word file Start *********************************************


# ************************* Upload Excel file Start *********************************************
# Function to safely convert exam years to integers
def safe_int(value, default=0):
    try:
        return int(value)
    except (ValueError, TypeError):
        return default  # Return a default value if conversion fails

def upload_file(request):
    if request.method == 'POST':
        form = UploadFileForm(request.POST, request.FILES)
        if form.is_valid():
            # Save the uploaded file temporarily
            file = request.FILES['file']
            fs = FileSystemStorage()
            filename = fs.save(file.name, file)
            uploaded_file_url = fs.url(filename)

            # Process the uploaded file
            data = pd.read_excel(fs.path(filename))

            # Replace NaN values with blank strings for text fields and 0 for numeric fields
            data = data.fillna({
                'marks': 0,
                'negative_marks': 0,
                'exam_year1': 0,
                'exam_year2': 0,
                'exam_year3': 0,
                'exam_year4': 0,
            }).fillna('')

            # Get the maximum question number from the database
            max_question_number = QuestionBank.objects.aggregate(Max('question_number'))['question_number__max']
            if max_question_number:
                start_number = int(max_question_number) + 1
            else:
                start_number = 1

            # Loop through the rows and create QuestionBank entries
            for _, row in data.iterrows():
                while QuestionBank.objects.filter(question_number=str(start_number)).exists():
                    start_number += 1

                # Set the type_of_question based on question_sub_type
                if row.get('question_sub_type') == 'true_and_false_type':
                    type_of_question = 'true_false'
                elif row.get('question_sub_type') == 'fill_in_the_blank_type':
                    type_of_question = 'fill_ups'
                else:
                    type_of_question = 'multiple_choice'

                # Handle multiple exam years using safe_int
                exam_years = [safe_int(row.get('exam_year1', 0)),
                              safe_int(row.get('exam_year2', 0)),
                              safe_int(row.get('exam_year3', 0)),
                              safe_int(row.get('exam_year4', 0))]
                
                # Filter out invalid years (like 0) and use the first valid one
                exam_year = next((year for year in exam_years if year != 0), 0)

                # Handle evergreen_index
                evergreen_index = row.get('evergreen_index', 5)  # Default value of 5
                if evergreen_index == '' or evergreen_index is None:
                    evergreen_index = 0  # Handle empty or None values
                else:
                    evergreen_index = int(evergreen_index)  # Ensure it's an integer

                # Get the ManyToMany fields for ExamName, Subject, Area, PartName, etc.
                exam_names = ExamName.objects.filter(name__in=[x.strip() for x in row.get('exam_name', '').split(',')])
                subjects = Subject.objects.filter(name__in=[x.strip() for x in row.get('subject_name', '').split(',')])
                areas = Area.objects.filter(name__in=[x.strip() for x in row.get('area_name', '').split(',')])
                parts = PartName.objects.filter(name__in=[x.strip() for x in row.get('part_name', '').split(',')])

                # Create the QuestionBank entry
                question = QuestionBank.objects.create(
                    question_number=str(start_number),
                    type_of_question=type_of_question,
                    exam_stage=row.get('exam_stage', ''),
                    exam_year=exam_year,  # Save the first valid exam year
                    language=row.get('language', ''),
                    script=row.get('script', ''),
                    evergreen_index=evergreen_index,  # Converted to integer
                    marks=float(row.get('exam_year1_marks', 0)),
                    negative_marks=float(row.get('exam_year1_negative_marks', 0)),
                    degree_of_difficulty=row.get('degree_of_difficulty', ''),
                    question_sub_type=row.get('question_sub_type', 'simple_type'),
                    question_part=row.get('question_part', ''),
                    question_part_first=row.get('question_part_first_part', ''),
                    assertion=row.get('assertion', ''),
                    reason=row.get('reason', ''),
                    list_1_name=row.get('list_1_name', ''),
                    list_2_name=row.get('list_2_name', ''),
                    list_1_row1=row.get('list_1_row1', ''),
                    list_2_row1=row.get('list_2_row1', ''),
                    list_1_row2=row.get('list_1_row2', ''),
                    list_2_row2=row.get('list_2_row2', ''),
                    list_1_row3=row.get('list_1_row3', ''),
                    list_2_row3=row.get('list_2_row3', ''),
                    list_1_row4=row.get('list_1_row4', ''),
                    list_2_row4=row.get('list_2_row4', ''),
                    list_1_row5=row.get('list_1_row5', ''),
                    list_2_row5=row.get('list_2_row5', ''),
                    list_1_row6=row.get('list_1_row6', ''),
                    list_2_row6=row.get('list_2_row6', ''),
                    list_1_row7=row.get('list_1_row7', ''),
                    list_2_row7=row.get('list_2_row7', ''),
                    list_1_row8=row.get('list_1_row8', ''),
                    list_2_row8=row.get('list_2_row8', ''),
                    question_part_third=row.get('question_part_third_part', ''),
                    answer_option_a=row.get('answer_option_a', ''),
                    answer_option_b=row.get('answer_option_b', ''),
                    answer_option_c=row.get('answer_option_c', ''),
                    answer_option_d=row.get('answer_option_d', ''),
                    correct_answer_choice=row.get('correct_answer_choice', ''),
                    correct_answer_description=row.get('correct_answer_description', ''),
                    created_by=request.user  # Assign the logged-in user
                )

                # ManyToMany relationships
                question.exam_name.set(exam_names)
                question.subject_name.set(subjects)
                question.area_name.set(areas)
                question.part_name.set(parts)

                start_number += 1

            messages.success(request, "File uploaded and data processed successfully!")
            return redirect('upload-file')
    else:
        form = UploadFileForm()

    return render(request, 'question_bank/upload.html', {'form': form})


def get_subjects(request):
    exam_id = request.GET.get('exam_id')
    subjects = Subject.objects.filter(exam_id=exam_id).values('id', 'name')
    return JsonResponse({'subjects': list(subjects)})

def get_areas(request):
    subject_id = request.GET.get('subject_id')
    areas = Area.objects.filter(subject_id=subject_id).values('id', 'name')
    return JsonResponse({'areas': list(areas)})

def get_parts(request):
    area_id = request.GET.get('area_id')
    parts = PartName.objects.filter(area_id=area_id).values('id', 'name')
    return JsonResponse({'parts': list(parts)})


def get_chapters(request):
    part_id = request.GET.get('part_id')
    chapeters = ChapterName.objects.filter(part_id=part_id).values('id', 'name')  # Typo: 'chapeters' should be 'chapters'
    return JsonResponse({'chapters': list(chapeters)})  # Also change to 'chapters'


def get_topics(request):
    chapter_id = request.GET.get('chapter_id')
    topics = TopicName.objects.filter(chapter_id=chapter_id).values('id', 'name')
    return JsonResponse({'topics': list(topics)})

# Get subjects for multiple exams
def get_subjects_list(request):
    exam_ids = request.GET.getlist('exam_ids[]')  # Capture multiple exam IDs
    subjects = Subject.objects.filter(exam_id__in=exam_ids).values('id', 'name')
    return JsonResponse({'subjects': list(subjects)})


def get_areas_list(request):
    subject_ids = request.GET.getlist('subject_ids[]')  # Capture multiple subject IDs
    areas = Area.objects.filter(subject_id__in=subject_ids).values('id', 'name')
    return JsonResponse({'areas': list(areas)})

def get_parts_list(request):
    area_ids = request.GET.getlist('area_ids[]')  # Capture multiple area IDs
    parts = PartName.objects.filter(area_id__in=area_ids).values('id', 'name')
    return JsonResponse({'parts': list(parts)})

def get_chapters_list(request):
    part_ids = request.GET.getlist('part_ids[]')  # Capture multiple part IDs
    chapters = ChapterName.objects.filter(part_id__in=part_ids).values('id', 'name')
    return JsonResponse({'chapters': list(chapters)})

def get_topics_list(request):
    chapter_ids = request.GET.getlist('chapter_ids[]')  # Capture multiple chapter IDs
    topics = TopicName.objects.filter(chapter_id__in=chapter_ids).values('id', 'name')
    return JsonResponse({'topics': list(topics)})





# ************************* Create Simple Type Question Start *********************************************
@login_required
def add_simple_type_question(request):
    if request.method == "POST":
        try:
            # Extract form data
            question_type = request.POST.get('questionType', 'simple_type')
            language = request.POST.get('language')
            script = request.POST.get('script')
            question_part_first = request.POST.get('question_part_first')
            answer_option_a = request.POST.get('answer_option_a')
            answer_option_b = request.POST.get('answer_option_b')
            answer_option_c = request.POST.get('answer_option_c')
            answer_option_d = request.POST.get('answer_option_d')
            correct_answer_choice = request.POST.get('correct_answer_choice')
            correct_answer_description = request.POST.get('correct_answer_description')

            # Get IDs for ManyToMany fields
            exam_name_ids = request.POST.getlist('exam_name[]')  # List of selected exam IDs
            subject_name_ids = request.POST.getlist('subject_name[]')
            area_name_ids = request.POST.getlist('area_name[]')
            part_name_ids = request.POST.getlist('part_name[]')
            chapter_name_ids = request.POST.getlist('chapter_name[]')
            topic_name_ids = request.POST.getlist('topic_name[]')  # List of selected topic IDs
            new_topic_name = request.POST.get('new_topic_name', None)  # Manually entered topic name

            # Extract other single values
            new_or_pyq = request.POST.get('new_or_pyq')
            exam_year = request.POST.get('exam_year', None)
            marks = float(request.POST.get('marks', 0.0))  # Convert to float
            negative_marks = float(request.POST.get('negative_marks', 0.0))  # Convert to float
            degree_of_difficulty = request.POST.get('degree_of_difficulty')
            evergreen_index = request.POST.get('evergreen_index')

            # Basic validation
            if not language or not script or not question_part_first:
                messages.error(request, "Please fill in all required fields.")
                return redirect('add-simple-type-question')

            # Check if the question type is 'pyq' to add the year, otherwise skip the year
            exam_year_value = exam_year if new_or_pyq == 'pyq' and exam_year else None

            # Handle topics: Add new topic if 'other' is selected
            if 'other' in topic_name_ids:
                topic_name_ids.remove('other')  # Remove 'other'
                if new_topic_name and chapter_name_ids:
                    selected_chapter = ChapterName.objects.get(id=chapter_name_ids[0])
                    new_topic = TopicName.objects.create(name=new_topic_name, chapter=selected_chapter)
                    topic_name_ids.append(new_topic.id)

            # Start a transaction to ensure atomicity
            with transaction.atomic():
                # Prepare data for ManyToMany relationships
                subjects = Subject.objects.filter(id__in=subject_name_ids)
                areas = Area.objects.filter(id__in=area_name_ids)
                parts = PartName.objects.filter(id__in=part_name_ids)
                chapters = ChapterName.objects.filter(id__in=chapter_name_ids)
                topics = TopicName.objects.filter(id__in=topic_name_ids)

                # Loop through each selected exam and create a question
                for exam_id in exam_name_ids:
                    last_question = QuestionBank.objects.order_by('-question_number').first()
                    next_question_number = last_question.question_number + 1 if last_question else 1

                    # Create the question
                    question = QuestionBank.objects.create(
                        question_number=next_question_number,
                        question_sub_type=question_type,
                        language=language,
                        script=script,
                        question_part=question_part_first,
                        answer_option_a=answer_option_a,
                        answer_option_b=answer_option_b,
                        answer_option_c=answer_option_c,
                        answer_option_d=answer_option_d,
                        correct_answer_choice=correct_answer_choice,
                        correct_answer_description=correct_answer_description,
                        exam_year=exam_year_value,
                        marks=marks,
                        negative_marks=negative_marks,
                        degree_of_difficulty=degree_of_difficulty,
                        evergreen_index=evergreen_index,
                        created_by=request.user  # Assign the logged-in user
                    )

                    # Assign ManyToMany fields
                    question.exam_name.set([ExamName.objects.get(id=exam_id)])
                    question.subject_name.set(subjects)
                    question.area_name.set(areas)
                    question.part_name.set(parts)
                    question.chapter_name.set(chapters)
                    question.topic_name.set(topics)

                    # Save the question
                    question.save()

            # Success message after all questions are added
            messages.success(request, "Questions added successfully!")
            return redirect('add-simple-type-question')

        except Exception as e:
            # Handle exceptions and rollback transaction
            messages.error(request, f"An error occurred: {e}")
            return redirect('add-simple-type-question')

    # If not POST, render the form with options for exams, subjects, etc.
    return render(request, 'question_bank/add_question/simple_type_form.html', {
        'exam_names': ExamName.objects.all(),
        'subjects': Subject.objects.all(),
        'topics': TopicName.objects.all()
    })


# ************************* Create Simple Type Question End *********************************************



# ************************* Create R and A Type Question Start *********************************************
@login_required
def add_r_and_a_type_question(request):
    # Fetch all the required data for dropdowns
    exam_names = ExamName.objects.all()
    topics = TopicName.objects.all()  # Fetch topics for the dropdown

    if request.method == 'POST':
        try:
            with transaction.atomic():  # Ensure all database operations are done atomically

                # Extract form data
                exam_name_ids = request.POST.getlist('exam_name[]')  # List of selected exam IDs
                subject_name_ids = request.POST.getlist('subject_name[]')
                area_name_ids = request.POST.getlist('area_name[]')
                part_name_ids = request.POST.getlist('part_name[]')
                chapter_name_ids = request.POST.getlist('chapter_name[]')
                topic_name_ids = request.POST.getlist('topic_name[]')
                new_topic_name = request.POST.get('new_topic_name', None)  # Manually entered topic name

                exam_year = request.POST.get('exam_year', None)
                if not exam_year:
                    exam_year = None

                # Topic handling: Check for existing topics or add new topic
                if 'other' in topic_name_ids:
                    # Remove 'other' from the list of topic_name_ids
                    topic_name_ids.remove('other')

                    # Check if a new topic was entered and add it to the database
                    if new_topic_name and chapter_name_ids:
                        # Ensure the new topic is linked to the selected chapter
                        selected_chapter = ChapterName.objects.get(id=chapter_name_ids[0])  # Assuming only one chapter is selected
                        new_topic, created = TopicName.objects.get_or_create(name=new_topic_name, chapter=selected_chapter)
                        # Add the new topic's ID to the list of selected topics
                        topic_name_ids.append(new_topic.id)

                # Loop through each selected exam and create a separate question
                for exam_id in exam_name_ids:
                    # Generate a unique question number (this can vary based on how you're handling question numbers)
                    last_question = QuestionBank.objects.order_by('-question_number').first()
                    next_question_number = last_question.question_number + 1 if last_question else 1

                    # Create a new question object for each exam
                    question = QuestionBank.objects.create(
                        question_number=next_question_number,  # Unique question number
                        question_sub_type=request.POST.get('questionType', 'r_and_a_type'),
                        question_part_first=request.POST.get('question_part_first', ''),
                        reason=request.POST.get('reason', ''),
                        assertion=request.POST.get('assertion', ''),
                        question_part_third=request.POST.get('question_part_third', ''),
                        correct_answer_choice=request.POST.get('correct_answer_choice', ''),
                        correct_answer_description=request.POST.get('correct_answer_description', ''),
                        exam_year=exam_year,  # Add year if PYQs, otherwise None
                        language=request.POST.get('language'),
                        script=request.POST.get('script'),
                        marks=float(request.POST.get('marks', 0.0)),  # Ensure marks is a float value
                        negative_marks=float(request.POST.get('negative_marks', 0.0)),  # Ensure negative marks is a float value
                        degree_of_difficulty=request.POST.get('degree_of_difficulty', ''),
                        evergreen_index=request.POST.get('evergreen_index', ''),
                        answer_option_a=request.POST.get('answer_option_a', ''),
                        answer_option_b=request.POST.get('answer_option_b', ''),
                        answer_option_c=request.POST.get('answer_option_c', ''),
                        answer_option_d=request.POST.get('answer_option_d', ''),
                        created_by=request.user  # Assign the logged-in user
                    )

                    # Add ManyToMany relationships (for each question, associate with the selected exam)
                    question.exam_name.set([ExamName.objects.get(id=exam_id)])  # Set only the current exam
                    question.subject_name.set(Subject.objects.filter(id__in=subject_name_ids))
                    question.area_name.set(Area.objects.filter(id__in=area_name_ids))
                    question.part_name.set(PartName.objects.filter(id__in=part_name_ids))
                    question.chapter_name.set(ChapterName.objects.filter(id__in=chapter_name_ids))
                    question.topic_name.set(TopicName.objects.filter(id__in=topic_name_ids))  # Set the topics (both existing and new)

                    # Save each question
                    question.save()

                # Redirect with a success message after creating all questions
                messages.success(request, 'R & A Type Question has been added successfully!')
                return redirect('add-r-and-a-type-question')

        except ExamName.DoesNotExist:
            messages.error(request, 'Selected exam does not exist.')
        except Subject.DoesNotExist:
            messages.error(request, 'Selected subject does not exist.')
        except Area.DoesNotExist:
            messages.error(request, 'Selected area does not exist.')
        except PartName.DoesNotExist:
            messages.error(request, 'Selected part does not exist.')
        except ChapterName.DoesNotExist:
            messages.error(request, 'Selected chapter does not exist.')
        except TopicName.DoesNotExist:
            messages.error(request, 'Selected topic does not exist.')
        except Exception as e:
            # General exception handler for any unexpected errors
            messages.error(request, f'An unexpected error occurred: {str(e)}')

    # Pass data to the form for dynamic dropdowns
    context = {
        'exam_names': exam_names,
        'topics': topics
    }

    return render(request, 'question_bank/add_question/r_and_a_type_form.html', context)


# ************************* Create R and A Type Question End *********************************************


# ************************* Create List-I Type Question Start *********************************************
def add_list_type_1_question(request):
    # Fetch all the required data for dropdowns
    exam_names = ExamName.objects.all()
    topics = TopicName.objects.all()  # Fetch topics for the dropdown

    if request.method == 'POST':
        try:
            with transaction.atomic():  # Ensure all database operations are done atomically
                # Extract form data
                exam_name_ids = request.POST.getlist('exam_name[]')
                subject_name_ids = request.POST.getlist('subject_name[]')
                area_name_ids = request.POST.getlist('area_name[]')
                part_name_ids = request.POST.getlist('part_name[]')
                chapter_name_ids = request.POST.getlist('chapter_name[]')
                topic_name_ids = request.POST.getlist('topic_name[]')
                new_topic_name = request.POST.get('new_topic_name', None)

                exam_year = request.POST.get('exam_year', None)
                if not exam_year:
                    exam_year = None

                # Handle topic name (can be selected from dropdown or manually added)
                if 'other' in topic_name_ids:
                    topic_name_ids.remove('other')
                    if new_topic_name and chapter_name_ids:
                        selected_chapter = ChapterName.objects.get(id=chapter_name_ids[0])
                        new_topic, created = TopicName.objects.get_or_create(name=new_topic_name, chapter=selected_chapter)
                        topic_name_ids.append(new_topic.id)

                # Loop through each selected exam and create a separate question
                for exam_id in exam_name_ids:
                    last_question = QuestionBank.objects.order_by('-question_number').first()
                    next_question_number = last_question.question_number + 1 if last_question else 1

                    question = QuestionBank.objects.create(
                        question_number=next_question_number,
                        question_sub_type=request.POST.get('questionType', 'list_type_1'),
                        question_part_first=request.POST.get('question_part_first', ''),
                        correct_answer_choice=request.POST.get('correct_answer_choice', ''),
                        correct_answer_description=request.POST.get('correct_answer_description', ''),
                        exam_year=exam_year,
                        language=request.POST.get('language'),
                        script=request.POST.get('script'),
                        marks=float(request.POST.get('marks', 0.0)),
                        negative_marks=float(request.POST.get('negative_marks', 0.0)),
                        degree_of_difficulty=request.POST.get('degree_of_difficulty', ''),
                        evergreen_index=request.POST.get('evergreen_index', ''),
                        list_1_row1=request.POST.get('list_1_row1', ''),
                        list_1_row2=request.POST.get('list_1_row2', ''),
                        list_1_row3=request.POST.get('list_1_row3', ''),
                        list_1_row4=request.POST.get('list_1_row4', ''),
                        list_1_row5=request.POST.get('list_1_row5', ''),
                        list_1_row6=request.POST.get('list_1_row6', ''),
                        list_1_row7=request.POST.get('list_1_row7', ''),
                        list_1_row8=request.POST.get('list_1_row8', ''),
                        question_part_third=request.POST.get('question_part_third', ''),
                        answer_option_a=request.POST.get('answer_option_a', ''),
                        answer_option_b=request.POST.get('answer_option_b', ''),
                        answer_option_c=request.POST.get('answer_option_c', ''),
                        answer_option_d=request.POST.get('answer_option_d', ''),
                        created_by=request.user
                    )

                    # Set ManyToMany fields
                    question.exam_name.set([ExamName.objects.get(id=exam_id)])
                    question.subject_name.set(Subject.objects.filter(id__in=subject_name_ids))
                    question.area_name.set(Area.objects.filter(id__in=area_name_ids))
                    question.part_name.set(PartName.objects.filter(id__in=part_name_ids))
                    question.chapter_name.set(ChapterName.objects.filter(id__in=chapter_name_ids))
                    question.topic_name.set(TopicName.objects.filter(id__in=topic_name_ids))

                    question.save()

                messages.success(request, 'List-I Type Question has been added successfully!')
                return redirect('add-list-type-1-question')

        except ExamName.DoesNotExist:
            messages.error(request, 'Selected exam does not exist.')
        except Subject.DoesNotExist:
            messages.error(request, 'Selected subject does not exist.')
        except Area.DoesNotExist:
            messages.error(request, 'Selected area does not exist.')
        except PartName.DoesNotExist:
            messages.error(request, 'Selected part does not exist.')
        except ChapterName.DoesNotExist:
            messages.error(request, 'Selected chapter does not exist.')
        except TopicName.DoesNotExist:
            messages.error(request, 'Selected topic does not exist.')
        except Exception as e:
            messages.error(request, f'An unexpected error occurred: {str(e)}')

    context = {
        'exam_names': exam_names,
        'topics': topics
    }

    return render(request, 'question_bank/add_question/list_type_1_form.html', context)



# ************************* Create List-I Type Question End *********************************************



# ************************* Create List-II Type Question Start *********************************************
def add_list_type_2_question(request):
    # Fetch all the required data for dropdowns
    exam_names = ExamName.objects.all()
    topics = TopicName.objects.all()  # Fetch topics for the dropdown

    if request.method == 'POST':
        try:
            with transaction.atomic():
                # Extract form data
                exam_name_ids = request.POST.getlist('exam_name[]')
                subject_name_ids = request.POST.getlist('subject_name[]')
                area_name_ids = request.POST.getlist('area_name[]')
                part_name_ids = request.POST.getlist('part_name[]')
                chapter_name_ids = request.POST.getlist('chapter_name[]')
                topic_name_ids = request.POST.getlist('topic_name[]')
                new_topic_name = request.POST.get('new_topic_name', None)

                exam_year = request.POST.get('exam_year', None)
                if not exam_year:
                    exam_year = None

                # Handle topic name (can be selected from dropdown or manually added)
                if 'other' in topic_name_ids:
                    topic_name_ids.remove('other')
                    if new_topic_name and chapter_name_ids:
                        selected_chapter = ChapterName.objects.get(id=chapter_name_ids[0])
                        new_topic, created = TopicName.objects.get_or_create(name=new_topic_name, chapter=selected_chapter)
                        topic_name_ids.append(new_topic.id)

                # Loop through each selected exam and create a separate question
                for exam_id in exam_name_ids:
                    last_question = QuestionBank.objects.order_by('-question_number').first()
                    next_question_number = last_question.question_number + 1 if last_question else 1

                    question = QuestionBank.objects.create(
                        question_number=next_question_number,
                        question_sub_type=request.POST.get('questionType', 'list_type_2'),
                        question_part_first=request.POST.get('question_part_first', ''),
                        correct_answer_choice=request.POST.get('correct_answer_choice', ''),
                        correct_answer_description=request.POST.get('correct_answer_description', ''),
                        exam_year=exam_year,
                        language=request.POST.get('language'),
                        script=request.POST.get('script'),
                        marks=float(request.POST.get('marks', 0.0)),
                        negative_marks=float(request.POST.get('negative_marks', 0.0)),
                        degree_of_difficulty=request.POST.get('degree_of_difficulty', ''),
                        evergreen_index=request.POST.get('evergreen_index', ''),
                        list_1_name=request.POST.get('list_1_name', ''),
                        list_2_name=request.POST.get('list_2_name', ''),
                        list_1_row1=request.POST.get('list_1_row1', ''),
                        list_2_row1=request.POST.get('list_2_row1', ''),
                        list_1_row2=request.POST.get('list_1_row2', ''),
                        list_2_row2=request.POST.get('list_2_row2', ''),
                        list_1_row3=request.POST.get('list_1_row3', ''),
                        list_2_row3=request.POST.get('list_2_row3', ''),
                        list_1_row4=request.POST.get('list_1_row4', ''),
                        list_2_row4=request.POST.get('list_2_row4', ''),
                        list_1_row5=request.POST.get('list_1_row5', ''),
                        list_2_row5=request.POST.get('list_2_row5', ''),
                        question_part_third=request.POST.get('question_part_third', ''),
                        answer_option_a=request.POST.get('answer_option_a', ''),
                        answer_option_b=request.POST.get('answer_option_b', ''),
                        answer_option_c=request.POST.get('answer_option_c', ''),
                        answer_option_d=request.POST.get('answer_option_d', ''),
                        created_by=request.user  # Assign the logged-in user
                    )

                    # Set ManyToMany fields
                    question.exam_name.set([ExamName.objects.get(id=exam_id)])
                    question.subject_name.set(Subject.objects.filter(id__in=subject_name_ids))
                    question.area_name.set(Area.objects.filter(id__in=area_name_ids))
                    question.part_name.set(PartName.objects.filter(id__in=part_name_ids))
                    question.chapter_name.set(ChapterName.objects.filter(id__in=chapter_name_ids))
                    question.topic_name.set(TopicName.objects.filter(id__in=topic_name_ids))

                    question.save()

                messages.success(request, 'List-II Type Question has been added successfully!')
                return redirect('add-list-type-2-question')

        except ExamName.DoesNotExist:
            messages.error(request, 'Selected exam does not exist.')
        except Subject.DoesNotExist:
            messages.error(request, 'Selected subject does not exist.')
        except Area.DoesNotExist:
            messages.error(request, 'Selected area does not exist.')
        except PartName.DoesNotExist:
            messages.error(request, 'Selected part does not exist.')
        except ChapterName.DoesNotExist:
            messages.error(request, 'Selected chapter does not exist.')
        except TopicName.DoesNotExist:
            messages.error(request, 'Selected topic does not exist.')
        except Exception as e:
            messages.error(request, f'An unexpected error occurred: {str(e)}')

    context = {
        'exam_names': exam_names,
        'topics': topics
    }

    return render(request, 'question_bank/add_question/list_type_2_form.html', context)


# ************************* Create List-II Type Question End *********************************************


# ************************* Create True and False Type Question Start *********************************************
def add_true_and_false_type_question(request):
    # Fetch all the required data for dropdowns
    exam_names = ExamName.objects.all()
    topics = TopicName.objects.all()  # Fetch topics for the dropdown

    if request.method == 'POST':
        try:
            with transaction.atomic():
                # Extract form data
                exam_name_ids = request.POST.getlist('exam_name[]')
                subject_name_ids = request.POST.getlist('subject_name[]')
                area_name_ids = request.POST.getlist('area_name[]')
                part_name_ids = request.POST.getlist('part_name[]')
                chapter_name_ids = request.POST.getlist('chapter_name[]')
                topic_name_ids = request.POST.getlist('topic_name[]')
                new_topic_name = request.POST.get('new_topic_name', None)

                exam_year = request.POST.get('exam_year', None)
                if not exam_year:
                    exam_year = None

                # Handle topic name (can be selected from dropdown or manually added)
                if 'other' in topic_name_ids:
                    topic_name_ids.remove('other')
                    if new_topic_name and chapter_name_ids:
                        selected_chapter = ChapterName.objects.get(id=chapter_name_ids[0])
                        new_topic, created = TopicName.objects.get_or_create(name=new_topic_name, chapter=selected_chapter)
                        topic_name_ids.append(new_topic.id)

                # Loop through each selected exam and create a separate question
                for exam_id in exam_name_ids:
                    last_question = QuestionBank.objects.order_by('-question_number').first()
                    next_question_number = last_question.question_number + 1 if last_question else 1

                    question = QuestionBank.objects.create(
                        question_number=next_question_number,
                        question_sub_type=request.POST.get('questionType', 'true_and_false_type'),
                        question_part=request.POST.get('question_part_first', ''),
                        correct_answer_choice=request.POST.get('correct_answer_choice', ''),
                        correct_answer_description=request.POST.get('correct_answer_description', ''),
                        exam_year=exam_year,
                        language=request.POST.get('language'),
                        script=request.POST.get('script'),
                        marks=float(request.POST.get('marks', 0.0)),
                        negative_marks=float(request.POST.get('negative_marks', 0.0)),
                        degree_of_difficulty=request.POST.get('degree_of_difficulty', ''),
                        evergreen_index=request.POST.get('evergreen_index', ''),
                        created_by=request.user,  # Assign the logged-in user
                        answer_option_a="True",
                        answer_option_b="False"
                    )

                    # Set ManyToMany fields
                    question.exam_name.set([ExamName.objects.get(id=exam_id)])
                    question.subject_name.set(Subject.objects.filter(id__in=subject_name_ids))
                    question.area_name.set(Area.objects.filter(id__in=area_name_ids))
                    question.part_name.set(PartName.objects.filter(id__in=part_name_ids))
                    question.chapter_name.set(ChapterName.objects.filter(id__in=chapter_name_ids))
                    question.topic_name.set(TopicName.objects.filter(id__in=topic_name_ids))

                    question.save()

                messages.success(request, 'True & False Type Question has been added successfully!')
                return redirect('add-true-and-false-type-question')

        except ExamName.DoesNotExist:
            messages.error(request, 'Selected exam does not exist.')
        except Subject.DoesNotExist:
            messages.error(request, 'Selected subject does not exist.')
        except Area.DoesNotExist:
            messages.error(request, 'Selected area does not exist.')
        except PartName.DoesNotExist:
            messages.error(request, 'Selected part does not exist.')
        except ChapterName.DoesNotExist:
            messages.error(request, 'Selected chapter does not exist.')
        except TopicName.DoesNotExist:
            messages.error(request, 'Selected topic does not exist.')
        except Exception as e:
            messages.error(request, f'An unexpected error occurred: {str(e)}')

    context = {
        'exam_names': exam_names,
        'topics': topics
    }

    return render(request, 'question_bank/add_question/true_false_type_form.html', context)

# ************************* Create True and False Type Question End *********************************************


# ************************* Create Fill in the Blank Type Question Start *********************************************
def add_fill_in_the_blank_question(request):
    # Fetch all required data for dropdowns
    exam_names = ExamName.objects.all()
    topics = TopicName.objects.all()  # Fetch topics for the dropdown

    if request.method == 'POST':
        try:
            with transaction.atomic():
                # Extract form data
                exam_name_ids = request.POST.getlist('exam_name[]')
                subject_name_ids = request.POST.getlist('subject_name[]')
                area_name_ids = request.POST.getlist('area_name[]')
                part_name_ids = request.POST.getlist('part_name[]')
                chapter_name_ids = request.POST.getlist('chapter_name[]')
                topic_name_ids = request.POST.getlist('topic_name[]')
                new_topic_name = request.POST.get('new_topic_name', None)

                exam_year = request.POST.get('exam_year', None)
                if not exam_year:
                    exam_year = None

                # Handle topic name (can be selected from dropdown or manually added)
                if 'other' in topic_name_ids:
                    topic_name_ids.remove('other')
                    if new_topic_name and chapter_name_ids:
                        selected_chapter = ChapterName.objects.get(id=chapter_name_ids[0])
                        new_topic, created = TopicName.objects.get_or_create(name=new_topic_name, chapter=selected_chapter)
                        topic_name_ids.append(new_topic.id)

                # Loop through each selected exam and create a separate question
                for exam_id in exam_name_ids:
                    last_question = QuestionBank.objects.order_by('-question_number').first()
                    next_question_number = last_question.question_number + 1 if last_question else 1

                    question = QuestionBank.objects.create(
                        question_number=next_question_number,
                        question_sub_type=request.POST.get('questionType', 'fill_in_the_blank_type'),
                        question_part=request.POST.get('question_part_first', ''),
                        correct_answer_choice=request.POST.get('correct_answer_choice', ''),
                        correct_answer_description=request.POST.get('correct_answer_description', ''),
                        exam_year=exam_year,
                        language=request.POST.get('language'),
                        script=request.POST.get('script'),
                        marks=float(request.POST.get('marks', 0.0)),
                        negative_marks=float(request.POST.get('negative_marks', 0.0)),
                        degree_of_difficulty=request.POST.get('degree_of_difficulty', ''),
                        evergreen_index=request.POST.get('evergreen_index', ''),
                        created_by=request.user  # Assign the logged-in user
                    )

                    # Set ManyToMany fields
                    question.exam_name.set([ExamName.objects.get(id=exam_id)])
                    question.subject_name.set(Subject.objects.filter(id__in=subject_name_ids))
                    question.area_name.set(Area.objects.filter(id__in=area_name_ids))
                    question.part_name.set(PartName.objects.filter(id__in=part_name_ids))
                    question.chapter_name.set(ChapterName.objects.filter(id__in=chapter_name_ids))
                    question.topic_name.set(TopicName.objects.filter(id__in=topic_name_ids))

                    question.save()

                messages.success(request, 'Fill in the Blank Question has been added successfully!')
                return redirect('add-fill-in-the-blank-question')

        except ExamName.DoesNotExist:
            messages.error(request, 'Selected exam does not exist.')
        except Subject.DoesNotExist:
            messages.error(request, 'Selected subject does not exist.')
        except Area.DoesNotExist:
            messages.error(request, 'Selected area does not exist.')
        except PartName.DoesNotExist:
            messages.error(request, 'Selected part does not exist.')
        except ChapterName.DoesNotExist:
            messages.error(request, 'Selected chapter does not exist.')
        except TopicName.DoesNotExist:
            messages.error(request, 'Selected topic does not exist.')
        except Exception as e:
            messages.error(request, f'An unexpected error occurred: {str(e)}')

    context = {
        'exam_names': exam_names,
        'topics': topics
    }

    return render(request, 'question_bank/add_question/fill_in_the_blank_form.html', context)



@login_required
def add_input_suggestion(request):
    # Fetch all required data for dropdowns
    exam_names = ExamName.objects.all()
    try:
        if request.method == 'POST':
            # Initialize the form with POST data and uploaded files
            form = InputSuggestionForm(request.POST, request.FILES)
            
            if form.is_valid():
                # Save the form but don’t commit to the database yet
                suggestion = form.save(commit=False)
                suggestion.created_by = request.user  # Set the logged-in user
                
                # Save the form to commit many-to-many relationships after initial save
                suggestion.save()

                # Handle file uploads for images
                if 'question_images' in request.FILES:
                    for image in request.FILES.getlist('question_images'):
                        InputSuggestionImage.objects.create(question=suggestion, image=image)

                # Handle file uploads for documents
                if 'question_documents' in request.FILES:
                    for document in request.FILES.getlist('question_documents'):
                        InputSuggestionDocument.objects.create(question=suggestion, document=document)

                # Display success message and redirect
                messages.success(request, 'Input Suggestion has been added successfully!')
                return redirect('input-suggestion-list')
        
    except Exception as e:
        return HttpResponse(f"Error: {str(e)}")

    context = {
        'exam_names': exam_names,
        'form': InputSuggestionForm()  # Passing the form with CKEditor widget enabled
    }
    
    return render(request, 'question_bank/add_input_suggestion.html', context)




def view_input_suggestion(request):
    # Fetch all DescriptiveTypeQuestion entries
    questions = InputSuggestion.objects.all()

    context = {
        'questions': questions
    }

    return render(request, 'question_bank/input_suggestion_list.html', context)


def question_blog_view(request, question_id):
    # Fetch the question using its ID
    question = get_object_or_404(InputSuggestion, id=question_id)

    context = {
        'question': question
    }
    
    return render(request, 'question_bank/view_input_suggestion.html', context)



@login_required
def view_questions(request):
    # Get questions only created by the logged-in user and prefetch the ManyToManyField data
    questions = QuestionBank.objects.filter(created_by=request.user).prefetch_related(
        'subject_name', 'area_name', 'part_name', 'chapter_name', 'topic_name'
    ).order_by('-created_at')

    context = {
        'questions': questions
    }
    return render(request, 'question_bank/add_question/view_questions.html', context)



def add_quote_idiom_phrase(request):
    if request.method == 'POST':
        # Get all form data
        type = request.POST.get('type')
        content = request.POST.get('content')
        meaning = request.POST.get('meaning', '')  # Optional meaning field
        author = request.POST.get('author', '')  # Optional author field
        exam_ids = request.POST.getlist('exam_name[]')  # Get multiple exam IDs
        subject_ids = request.POST.getlist('subject_name[]')  # Get multiple subject IDs
        area_ids = request.POST.getlist('area_name[]')  # Get multiple area IDs
        part_ids = request.POST.getlist('part_name[]')  # Get multiple part IDs
        chapter_ids = request.POST.getlist('chapter_name[]')  # Get multiple chapter IDs
        topic_ids = request.POST.getlist('topic_name[]')  # Get multiple topic IDs
        new_topic_name = request.POST.get('new_topic_name', '')  # New topic name if manually entered

        # Remove 'other' from topic_ids, as 'other' is not a valid ID
        if 'other' in topic_ids:
            topic_ids.remove('other')

        # Create the QuoteIdiomPhrase instance
        new_entry = QuoteIdiomPhrase.objects.create(
            type=type,
            content=content,
            meaning=meaning if type in ['idiom', 'phrase'] else '',  # Add meaning only for idioms or phrases
            author=author,  # Add author if provided
            created_by=request.user
        )

        # Handle ManyToMany fields after creating the instance
        if exam_ids:
            exams = ExamName.objects.filter(id__in=exam_ids)
            new_entry.exams.set(exams)

        if subject_ids:
            subjects = Subject.objects.filter(id__in=subject_ids)
            new_entry.subjects.set(subjects)

        if area_ids:
            areas = Area.objects.filter(id__in=area_ids)
            new_entry.areas.set(areas)

        if part_ids:
            parts = PartName.objects.filter(id__in=part_ids)
            new_entry.parts.set(parts)

        if chapter_ids:
            chapters = ChapterName.objects.filter(id__in=chapter_ids)
            new_entry.chapters.set(chapters)

        # Handle manually entered topic with association to the selected chapter
        if new_topic_name and chapter_ids:
            # Associate the manually entered topic with the first selected chapter
            related_chapter = ChapterName.objects.filter(id=chapter_ids[0]).first()
            new_topic = TopicName.objects.create(name=new_topic_name, chapter=related_chapter)
            new_entry.topics.add(new_topic)

        # Handle selected topics (without 'other')
        if topic_ids:
            selected_topics = TopicName.objects.filter(id__in=topic_ids)
            new_entry.topics.add(*selected_topics)

        # Display success message and redirect
        messages.success(request, 'Your Quote, Idiom, or Phrase has been added successfully!')
        return redirect('add_quote_idiom_phrase')

    # Fetch all exam names for the form
    exam_names = ExamName.objects.all()
    return render(request, 'question_bank/add_quote_idiom_phrase.html', {'exam_names': exam_names})





def quotes_idioms_phrases_view(request):
    # Fetch all quotes, idioms, and phrases from the database
    quotes_idioms_phrases = QuoteIdiomPhrase.objects.all().order_by('-created_at')
    
    # Render the template with the fetched data
    return render(request, 'question_bank/quotes_idioms_phrases.html', {
        'quotes_idioms_phrases': quotes_idioms_phrases
    })





def analytics_dashboard(request):
    # Filter and annotate data for users with questions, suggestions, or quote/phrase/idiom entries
    users_data = User.objects.annotate(
        user_name=Concat('first_name', Value(' '), 'last_name'),
        total_input_suggestions=Count('inputsuggestion', filter=Q(inputsuggestion__isnull=False)),
        total_questions=Count('questionbank', filter=Q(questionbank__isnull=False)),
        simple_type_count=Count('questionbank', filter=Q(questionbank__question_sub_type='simple_type')),
        list1_type_count=Count('questionbank', filter=Q(questionbank__question_sub_type='list_type_1')),
        list2_type_count=Count('questionbank', filter=Q(questionbank__question_sub_type='list_type_2')),
        ra_type_count=Count('questionbank', filter=Q(questionbank__question_sub_type='r_and_a_type')),
        true_false_type_count=Count('questionbank', filter=Q(questionbank__question_sub_type='true_and_false_type')),
        fill_blank_type_count=Count('questionbank', filter=Q(questionbank__question_sub_type='fill_in_the_blank_type')),
        phrases_uploaded=Count('quoteidiomphrase', filter=Q(quoteidiomphrase__type='phrase')),
        idioms_uploaded=Count('quoteidiomphrase', filter=Q(quoteidiomphrase__type='idiom')),
        quotes_uploaded=Count('quoteidiomphrase', filter=Q(quoteidiomphrase__type='quote')),
    ).filter(
        Q(total_input_suggestions__gt=0) | 
        Q(total_questions__gt=0) | 
        Q(phrases_uploaded__gt=0) |
        Q(idioms_uploaded__gt=0) | 
        Q(quotes_uploaded__gt=0)
    ).values_list(
        'user_name', 'email', 'total_input_suggestions', 'total_questions', 
        'simple_type_count', 'list1_type_count', 'list2_type_count',
        'ra_type_count', 'true_false_type_count', 'fill_blank_type_count',
        'phrases_uploaded', 'idioms_uploaded', 'quotes_uploaded',
        named=True
    )
    # Prepare data for charts
    users = [user.user_name for user in users_data]
    emails = [user.email for user in users_data]
    suggestion_counts = [user.total_input_suggestions for user in users_data]
    question_counts = [user.total_questions for user in users_data]
    simple_counts = [user.simple_type_count for user in users_data]
    list1_counts = [user.list1_type_count for user in users_data]
    list2_counts = [user.list2_type_count for user in users_data]
    ra_counts = [user.ra_type_count for user in users_data]
    true_false_counts = [user.true_false_type_count for user in users_data]
    fill_blank_counts = [user.fill_blank_type_count for user in users_data]
    phrase_counts = [user.phrases_uploaded for user in users_data]
    idiom_counts = [user.idioms_uploaded for user in users_data]
    quote_counts = [user.quotes_uploaded for user in users_data]

    context = {
        'users': json.dumps(users),
        'emails': json.dumps(emails),
        'suggestion_counts': json.dumps(suggestion_counts),
        'question_counts': json.dumps(question_counts),
        'simple_counts': json.dumps(simple_counts),
        'list1_counts': json.dumps(list1_counts),
        'list2_counts': json.dumps(list2_counts),
        'ra_counts': json.dumps(ra_counts),
        'true_false_counts': json.dumps(true_false_counts),
        'fill_blank_counts': json.dumps(fill_blank_counts),
        'phrase_counts': json.dumps(phrase_counts),
        'idiom_counts': json.dumps(idiom_counts),
        'quote_counts': json.dumps(quote_counts),
    }

    return render(request, 'question_bank/analytics_dashboard.html', context)


@login_required
def new_dashboard_view(request):
    today = timezone.now().date()
    this_week_start = today - timedelta(days=7)  # Last 7 days
    
    # Date range for earlier week (8–14 days ago)
    earlier_week_start, earlier_week_end = today - timedelta(days=14), today - timedelta(days=7)

    user_is_admin = request.user.user_type == 'admin'

    # Prepare data for this week
    this_week_report_data = []
    earlier_week_report_data = []

    # Fetch users based on user type
    users = User.objects.all() if user_is_admin else User.objects.filter(id=request.user.id)

    # Collect This Week's Reports (last 7 days)
    for user in users:
        total_questions = QuestionBank.objects.filter(created_by=user, created_at__gte=this_week_start).count()
        total_phrases = QuoteIdiomPhrase.objects.filter(created_by=user, created_at__gte=this_week_start).count()
        total_suggestions = InputSuggestion.objects.filter(created_by=user, created_at__gte=this_week_start).count()

        simple_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=this_week_start, question_sub_type='simple_type').count()
        list_1_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=this_week_start, question_sub_type='list_type_1').count()
        list_2_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=this_week_start, question_sub_type='list_type_2').count()
        ra_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=this_week_start, question_sub_type='r_and_a_type').count()
        true_false_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=this_week_start, question_sub_type='true_and_false_type').count()
        fill_blank_count = QuestionBank.objects.filter(created_by=user, created_at__gte=this_week_start, question_sub_type='fill_in_the_blank_type').count()

        this_week_report_data.append({
            'user': user,
            'total_questions': total_questions,
            'total_phrases': total_phrases,
            'total_suggestions': total_suggestions,
            'simple_type_count': simple_type_count,
            'list_1_type_count': list_1_type_count,
            'list_2_type_count': list_2_type_count,
            'ra_type_count': ra_type_count,
            'true_false_type_count': true_false_type_count,
            'fill_blank_count': fill_blank_count,
        })

    # Collect Earlier Week's Reports (8–14 days ago)
    for user in users:
        total_questions = QuestionBank.objects.filter(created_by=user, created_at__gte=earlier_week_start, created_at__lt=earlier_week_end).count()
        total_phrases = QuoteIdiomPhrase.objects.filter(created_by=user, created_at__gte=earlier_week_start, created_at__lt=earlier_week_end).count()
        total_suggestions = InputSuggestion.objects.filter(created_by=user, created_at__gte=earlier_week_start, created_at__lt=earlier_week_end).count()

        simple_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=earlier_week_start, created_at__lt=earlier_week_end, question_sub_type='simple_type').count()
        list_1_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=earlier_week_start, created_at__lt=earlier_week_end, question_sub_type='list_type_1').count()
        list_2_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=earlier_week_start, created_at__lt=earlier_week_end, question_sub_type='list_type_2').count()
        ra_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=earlier_week_start, created_at__lt=earlier_week_end, question_sub_type='r_and_a_type').count()
        true_false_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=earlier_week_start, created_at__lt=earlier_week_end, question_sub_type='true_and_false_type').count()
        fill_blank_count = QuestionBank.objects.filter(created_by=user, created_at__gte=earlier_week_start, created_at__lt=earlier_week_end, question_sub_type='fill_in_the_blank_type').count()

        earlier_week_report_data.append({
            'user': user,
            'total_questions': total_questions,
            'total_phrases': total_phrases,
            'total_suggestions': total_suggestions,
            'simple_type_count': simple_type_count,
            'list_1_type_count': list_1_type_count,
            'list_2_type_count': list_2_type_count,
            'ra_type_count': ra_type_count,
            'true_false_type_count': true_false_type_count,
            'fill_blank_count': fill_blank_count,
        })

    # Get the last generated dates
    this_week_generated_date = Report.objects.filter(report_type='this_week').order_by('-report_date').first()
    earlier_week_generated_date = Report.objects.filter(report_type='earlier').order_by('-report_date').first()

    context = {
        'this_week_report_data': this_week_report_data,
        'earlier_week_report_data': earlier_week_report_data,
        'this_week_generated_date': this_week_generated_date.report_date if this_week_generated_date else None,
        'earlier_week_generated_date': earlier_week_generated_date.report_date if earlier_week_generated_date else None,
    }
    return render(request, 'question_bank/dashboard.html', context)





@login_required
def generate_this_week_csv(request):
    today = timezone.now().date()
    this_week_start = today - timedelta(days=7)

    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="this_week_report.csv"'
    writer = csv.writer(response)

    writer.writerow(['Sr. No.', 'Email', 'User Name', 'No. of Questions Uploaded', 'No. of Phrases/Idioms/Quotes', 'No. of Input Suggestions',
                     'Simple Type', 'List I Type', 'List II Type', 'R & A Type', 'True & False', 'Fill in the Blank'])

    users = User.objects.all() if request.user.user_type == 'admin' else User.objects.filter(id=request.user.id)

    for idx, user in enumerate(users, 1):
        total_questions = QuestionBank.objects.filter(created_by=user, created_at__gte=this_week_start).count()
        total_phrases = QuoteIdiomPhrase.objects.filter(created_by=user, created_at__gte=this_week_start).count()
        total_suggestions = InputSuggestion.objects.filter(created_by=user, created_at__gte=this_week_start).count()

        simple_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=this_week_start, question_sub_type='simple_type').count()
        list_1_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=this_week_start, question_sub_type='list_type_1').count()
        list_2_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=this_week_start, question_sub_type='list_type_2').count()
        ra_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=this_week_start, question_sub_type='r_and_a_type').count()
        true_false_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=this_week_start, question_sub_type='true_and_false_type').count()
        fill_blank_count = QuestionBank.objects.filter(created_by=user, created_at__gte=this_week_start, question_sub_type='fill_in_the_blank_type').count()

        writer.writerow([idx, user.email, f"{user.first_name} {user.last_name}", total_questions, total_phrases, total_suggestions,
                         simple_type_count, list_1_type_count, list_2_type_count, ra_type_count, true_false_type_count, fill_blank_count])

        # Save the report in the Report model
        Report.objects.create(
            report_type='this_week',
            report_date=today,
            created_by=user,
            total_questions=total_questions,
            total_phrases=total_phrases,
            total_suggestions=total_suggestions,
            simple_type_count=simple_type_count,
            list_1_type_count=list_1_type_count,
            list_2_type_count=list_2_type_count,
            ra_type_count=ra_type_count,
            true_false_type_count=true_false_type_count,
            fill_blank_count=fill_blank_count,
        )

    return response



@login_required
def generate_earlier_week_csv(request):
    earlier_week_start, earlier_week_end = timezone.now() - timedelta(days=14), timezone.now() - timedelta(days=7)

    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="earlier_week_report.csv"'
    writer = csv.writer(response)

    writer.writerow(['Sr. No.', 'Email', 'User Name', 'No. of Questions Uploaded', 'No. of Phrases/Idioms/Quotes', 'No. of Input Suggestions',
                     'Simple Type', 'List I Type', 'List II Type', 'R & A Type', 'True & False', 'Fill in the Blank'])

    users = User.objects.all() if request.user.user_type == 'admin' else User.objects.filter(id=request.user.id)

    for idx, user in enumerate(users, 1):
        total_questions = QuestionBank.objects.filter(created_by=user, created_at__gte=earlier_week_start, created_at__lt=earlier_week_end).count()
        total_phrases = QuoteIdiomPhrase.objects.filter(created_by=user, created_at__gte=earlier_week_start, created_at__lt=earlier_week_end).count()
        total_suggestions = InputSuggestion.objects.filter(created_by=user, created_at__gte=earlier_week_start, created_at__lt=earlier_week_end).count()

        simple_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=earlier_week_start, created_at__lt=earlier_week_end, question_sub_type='simple_type').count()
        list_1_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=earlier_week_start, created_at__lt=earlier_week_end, question_sub_type='list_type_1').count()
        list_2_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=earlier_week_start, created_at__lt=earlier_week_end, question_sub_type='list_type_2').count()
        ra_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=earlier_week_start, created_at__lt=earlier_week_end, question_sub_type='r_and_a_type').count()
        true_false_type_count = QuestionBank.objects.filter(created_by=user, created_at__gte=earlier_week_start, created_at__lt=earlier_week_end, question_sub_type='true_and_false_type').count()
        fill_blank_count = QuestionBank.objects.filter(created_by=user, created_at__gte=earlier_week_start, created_at__lt=earlier_week_end, question_sub_type='fill_in_the_blank_type').count()

        writer.writerow([idx, user.email, f"{user.first_name} {user.last_name}", total_questions, total_phrases, total_suggestions,
                         simple_type_count, list_1_type_count, list_2_type_count, ra_type_count, true_false_type_count, fill_blank_count])

        # Save the report in the Report model
        Report.objects.create(
            report_type='earlier',
            report_date=timezone.now(),
            created_by=user,
            total_questions=total_questions,
            total_phrases=total_phrases,
            total_suggestions=total_suggestions,
            simple_type_count=simple_type_count,
            list_1_type_count=list_1_type_count,
            list_2_type_count=list_2_type_count,
            ra_type_count=ra_type_count,
            true_false_type_count=true_false_type_count,
            fill_blank_count=fill_blank_count,
        )

    return response
